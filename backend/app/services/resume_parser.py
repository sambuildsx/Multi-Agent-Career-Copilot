import os
import pdfplumber
from docx import Document


class PDFService:
    """Extracts raw text from resume files. Despite the name, this now
    handles both PDF and Word documents, since the upload UI accepts both."""

    @staticmethod
    def extract_text(file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return PDFService._extract_from_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return PDFService._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported resume file type: '{ext}'. Expected .pdf, .docx, or .doc.")

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Word resumes often put content in tables (skills tables, etc.) —
        # paragraphs alone would silently miss that content.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        return "\n".join(paragraphs)